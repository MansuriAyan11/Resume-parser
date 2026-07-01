"""DOCX text extraction using python-docx."""

from __future__ import annotations

import io
import logging
from pathlib import Path

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from resume_parser.exceptions import TextExtractionError
from resume_parser.extractors.base import BaseTextExtractor
from resume_parser.utils.text_utils import normalize_whitespace

logger = logging.getLogger(__name__)


class DocxTextExtractor(BaseTextExtractor):
    """Extract text from DOCX resume files."""

    def extract_from_path(self, file_path: Path) -> str:
        logger.debug("Extracting text from DOCX: %s", file_path)
        try:
            document = Document(str(file_path))
            return self._extract_from_document(document)
        except PackageNotFoundError as exc:
            logger.exception("Invalid DOCX package: %s", file_path)
            raise TextExtractionError(
                f"Invalid or corrupted DOCX file '{file_path.name}'", cause=exc
            ) from exc
        except Exception as exc:
            logger.exception("DOCX extraction failed for %s", file_path)
            raise TextExtractionError(
                f"Could not extract text from DOCX '{file_path.name}'", cause=exc
            ) from exc

    def extract_from_bytes(self, content: bytes, extension: str) -> str:
        logger.debug("Extracting text from DOCX bytes (%d bytes)", len(content))
        try:
            document = Document(io.BytesIO(content))
            return self._extract_from_document(document)
        except Exception as exc:
            logger.exception("DOCX byte extraction failed")
            raise TextExtractionError(
                "Could not extract text from DOCX bytes", cause=exc
            ) from exc

    def _extract_from_document(self, document: Document) -> str:
        paragraphs: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                paragraphs.append(text)

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        combined = "\n".join(paragraphs)
        try:
            return normalize_whitespace(self._validate_text(combined))
        except ValueError as exc:
            raise TextExtractionError("DOCX contains no extractable text", cause=exc) from exc
